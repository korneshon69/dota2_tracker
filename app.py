import csv
import io
from datetime import datetime
from functools import wraps

from flask import (Flask, render_template, redirect, url_for, flash, request,
                   abort, Response, send_file)
from flask_login import (LoginManager, login_user, logout_user, login_required,
                         current_user)
from flask_wtf.csrf import CSRFProtect
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import Config
from models import db, User, Hero, Match, HeroRole, MatchType
from forms import (LoginForm, RegisterForm, MatchForm, HeroForm,
                   HeroRoleForm, MatchTypeForm, UserEditForm)

app = Flask(__name__)
app.config.from_object(Config)

db.init_app(app)
csrf = CSRFProtect(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Пожалуйста, войдите в систему.'
login_manager.login_message_category = 'warning'


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# ---------- Декоратор для админа ----------

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            abort(403)
        return f(*args, **kwargs)
    return decorated_function


# ---------- Контекстный процессор ----------

@app.context_processor
def inject_now():
    return {'now': datetime.today}


# ---------- Обработчики ошибок ----------

@app.errorhandler(403)
def forbidden(e):
    return render_template('base.html', error_code=403,
                           error_message='Доступ запрещён'), 403


@app.errorhandler(404)
def not_found(e):
    return render_template('base.html', error_code=404,
                           error_message='Страница не найдена'), 404


# ==========================================
#          ОБЩИЕ МАРШРУТЫ
# ==========================================

@app.route('/')
def index():
    total_users = User.query.count()
    total_matches = Match.query.count()
    total_heroes = Hero.query.count()
    return render_template('index.html',
                           total_users=total_users,
                           total_matches=total_matches,
                           total_heroes=total_heroes)


# ---------- Аутентификация ----------

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = RegisterForm()
    if form.validate_on_submit():
        user = User(username=form.username.data, email=form.email.data)
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('Регистрация прошла успешно! Теперь войдите.', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user is None or not user.check_password(form.password.data):
            flash('Неверное имя пользователя или пароль.', 'danger')
            return redirect(url_for('login'))
        if not user.is_active_user:
            flash('Ваш аккаунт деактивирован.', 'danger')
            return redirect(url_for('login'))
        login_user(user, remember=form.remember.data)
        flash(f'Добро пожаловать, {user.username}!', 'success')
        next_page = request.args.get('next')
        return redirect(next_page or url_for('index'))
    return render_template('login.html', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Вы вышли из системы.', 'info')
    return redirect(url_for('index'))


# ---------- Герои (пользовательский просмотр) ----------

@app.route('/heroes')
def heroes():
    # Фильтры
    attribute = request.args.get('attribute', '')
    role_id = request.args.get('role_id', '', type=str)
    search = request.args.get('search', '')

    query = Hero.query

    if attribute:
        query = query.filter_by(attribute=attribute)
    if role_id:
        query = query.filter_by(role_id=int(role_id))
    if search:
        query = query.filter(Hero.name.ilike(f'%{search}%'))

    heroes_list = query.order_by(Hero.name).all()
    roles = HeroRole.query.order_by(HeroRole.name).all()

    return render_template('heroes.html', heroes=heroes_list, roles=roles,
                           attribute=attribute, role_id=role_id, search=search)


@app.route('/heroes/<int:hero_id>')
def hero_detail(hero_id):
    hero = Hero.query.get_or_404(hero_id)
    global_winrate = hero.get_global_winrate()
    total_picks = hero.matches.count()

    user_winrate = None
    user_matches = None
    if current_user.is_authenticated:
        user_winrate = hero.get_winrate_for_user(current_user.id)
        user_matches = hero.matches.filter_by(user_id=current_user.id).order_by(
            Match.date_played.desc()).limit(10).all()

    return render_template('hero_detail.html', hero=hero,
                           global_winrate=global_winrate,
                           total_picks=total_picks,
                           user_winrate=user_winrate,
                           user_matches=user_matches)


# ---------- Матчи (пользователь) ----------

@app.route('/matches')
@login_required
def matches():
    page = request.args.get('page', 1, type=int)

    # Фильтры
    hero_id = request.args.get('hero_id', '', type=str)
    result = request.args.get('result', '')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    team_role = request.args.get('team_role', '')

    query = Match.query.filter_by(user_id=current_user.id)

    if hero_id:
        query = query.filter_by(hero_id=int(hero_id))
    if result:
        query = query.filter_by(result=result)
    if team_role:
        query = query.filter_by(team_role=team_role)
    if date_from:
        try:
            df = datetime.strptime(date_from, '%Y-%m-%d')
            query = query.filter(Match.date_played >= df)
        except ValueError:
            pass
    if date_to:
        try:
            dt = datetime.strptime(date_to, '%Y-%m-%d')
            dt = dt.replace(hour=23, minute=59, second=59)
            query = query.filter(Match.date_played <= dt)
        except ValueError:
            pass

    pagination = query.order_by(Match.date_played.desc()).paginate(
        page=page, per_page=app.config['ITEMS_PER_PAGE'], error_out=False)

    heroes_list = Hero.query.order_by(Hero.name).all()

    return render_template('matches.html', pagination=pagination,
                           heroes=heroes_list, hero_id=hero_id,
                           result=result, date_from=date_from,
                           date_to=date_to, team_role=team_role)


@app.route('/matches/add', methods=['GET', 'POST'])
@login_required
def add_match():
    form = MatchForm()
    form.hero_id.choices = [(h.id, h.name) for h in Hero.query.order_by(Hero.name).all()]
    form.match_type_id.choices = [(t.id, t.name) for t in MatchType.query.order_by(MatchType.name).all()]

    if form.validate_on_submit():
        match = Match(
            user_id=current_user.id,
            hero_id=form.hero_id.data,
            match_type_id=form.match_type_id.data,
            date_played=form.date_played.data,
            duration_minutes=form.duration_minutes.data,
            result=form.result.data,
            kills=form.kills.data,
            deaths=form.deaths.data,
            assists=form.assists.data,
            team_role=form.team_role.data,
            gpm=form.gpm.data or 0,
            xpm=form.xpm.data or 0,
            notes=form.notes.data or ''
        )
        db.session.add(match)
        db.session.commit()
        flash('Матч успешно добавлен!', 'success')
        return redirect(url_for('matches'))

    # Установим дату по умолчанию
    if not form.date_played.data:
        form.date_played.data = datetime.utcnow()

    return render_template('add_match.html', form=form, title='Добавить матч')


@app.route('/matches/<int:match_id>')
@login_required
def match_detail(match_id):
    match = Match.query.get_or_404(match_id)
    if match.user_id != current_user.id and not current_user.is_admin:
        abort(403)
    return render_template('match_detail.html', match=match)


@app.route('/matches/<int:match_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_match(match_id):
    match = Match.query.get_or_404(match_id)
    if match.user_id != current_user.id and not current_user.is_admin:
        abort(403)

    form = MatchForm(obj=match)
    form.hero_id.choices = [(h.id, h.name) for h in Hero.query.order_by(Hero.name).all()]
    form.match_type_id.choices = [(t.id, t.name) for t in MatchType.query.order_by(MatchType.name).all()]

    if form.validate_on_submit():
        match.hero_id = form.hero_id.data
        match.match_type_id = form.match_type_id.data
        match.date_played = form.date_played.data
        match.duration_minutes = form.duration_minutes.data
        match.result = form.result.data
        match.kills = form.kills.data
        match.deaths = form.deaths.data
        match.assists = form.assists.data
        match.team_role = form.team_role.data
        match.gpm = form.gpm.data or 0
        match.xpm = form.xpm.data or 0
        match.notes = form.notes.data or ''
        db.session.commit()
        flash('Матч обновлён!', 'success')
        return redirect(url_for('match_detail', match_id=match.id))

    return render_template('edit_match.html', form=form, match=match,
                           title='Редактировать матч')


@app.route('/matches/<int:match_id>/delete', methods=['POST'])
@login_required
def delete_match(match_id):
    match = Match.query.get_or_404(match_id)
    if match.user_id != current_user.id and not current_user.is_admin:
        abort(403)
    db.session.delete(match)
    db.session.commit()
    flash('Матч удалён.', 'info')
    return redirect(url_for('matches'))


# ---------- Статистика пользователя ----------

@app.route('/stats')
@login_required
def stats():
    total_matches = current_user.matches.count()
    wins = current_user.matches.filter_by(result='win').count()
    losses = total_matches - wins
    winrate = current_user.get_winrate()
    avg_kda = current_user.get_avg_kda()
    best_hero = current_user.get_best_hero()

    # Средние показатели
    avg_stats = db.session.query(
        db.func.avg(Match.kills).label('avg_kills'),
        db.func.avg(Match.deaths).label('avg_deaths'),
        db.func.avg(Match.assists).label('avg_assists'),
        db.func.avg(Match.duration_minutes).label('avg_duration'),
        db.func.avg(Match.gpm).label('avg_gpm'),
        db.func.avg(Match.xpm).label('avg_xpm')
    ).filter_by(user_id=current_user.id).first()

    # Статистика по героям
    hero_stats = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins'),
        db.func.avg(Match.kills).label('avg_k'),
        db.func.avg(Match.deaths).label('avg_d'),
        db.func.avg(Match.assists).label('avg_a')
    ).join(Match).filter(
        Match.user_id == current_user.id
    ).group_by(Hero.id).order_by(db.text('games DESC')).all()

    # Статистика по ролям
    role_stats = db.session.query(
        Match.team_role,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).filter_by(user_id=current_user.id).group_by(
        Match.team_role
    ).all()

    return render_template('stats.html',
                           total_matches=total_matches, wins=wins, losses=losses,
                           winrate=winrate, avg_kda=avg_kda, best_hero=best_hero,
                           avg_stats=avg_stats, hero_stats=hero_stats,
                           role_stats=role_stats)


# ---------- Отчёты и экспорт ----------

@app.route('/reports')
@login_required
def reports():
    # Топ героев по винрейту (минимум 3 игры)
    top_heroes = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).join(Match).filter(
        Match.user_id == current_user.id
    ).group_by(Hero.id).having(
        db.func.count(Match.id) >= 1
    ).all()

    top_heroes_data = []
    for h in top_heroes:
        wr = round(h.wins / h.games * 100, 1) if h.games > 0 else 0
        top_heroes_data.append({
            'name': h.name, 'games': h.games,
            'wins': h.wins, 'winrate': wr
        })
    top_heroes_data.sort(key=lambda x: x['winrate'], reverse=True)

    # Активность по месяцам
    monthly = db.session.query(
        db.func.strftime('%Y-%m', Match.date_played).label('month'),
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).filter_by(user_id=current_user.id).group_by('month').order_by('month').all()

    return render_template('reports.html',
                           top_heroes=top_heroes_data,
                           monthly=monthly)


@app.route('/export/matches')
@login_required
def export_matches():
    matches_list = Match.query.filter_by(user_id=current_user.id).order_by(
        Match.date_played.desc()).all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Дата', 'Герой', 'Тип матча', 'Результат', 'K', 'D', 'A',
                     'KDA', 'Длительность (мин)', 'Роль', 'GPM', 'XPM', 'Заметки'])

    for m in matches_list:
        writer.writerow([
            m.date_played.strftime('%Y-%m-%d %H:%M'),
            m.hero.name,
            m.match_type_ref.name if m.match_type_ref else '',
            'Победа' if m.result == 'win' else 'Поражение',
            m.kills, m.deaths, m.assists,
            m.kda_ratio,
            m.duration_minutes,
            m.team_role,
            m.gpm, m.xpm,
            m.notes
        ])

    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment;filename=dota2_matches_export.csv'}
    )


@app.route('/export/stats')
@login_required
def export_stats():
    hero_stats = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins'),
        db.func.avg(Match.kills).label('avg_k'),
        db.func.avg(Match.deaths).label('avg_d'),
        db.func.avg(Match.assists).label('avg_a')
    ).join(Match).filter(
        Match.user_id == current_user.id
    ).group_by(Hero.id).order_by(db.text('games DESC')).all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Герой', 'Игры', 'Победы', 'Винрейт %', 'Ср. убийства',
                     'Ср. смерти', 'Ср. помощь'])

    for h in hero_stats:
        wr = round(h.wins / h.games * 100, 1) if h.games > 0 else 0
        writer.writerow([
            h.name, h.games, h.wins, wr,
            round(h.avg_k, 1), round(h.avg_d, 1), round(h.avg_a, 1)
        ])

    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment;filename=dota2_stats_export.csv'}
    )

# ---------- Экспорт в Word (пользователь) ----------

@app.route('/export/matches/word')
@login_required
def export_matches_word():
    matches_list = Match.query.filter_by(user_id=current_user.id).order_by(
        Match.date_played.desc()).all()

    doc = create_word_document(f'История матчей — {current_user.username}')

    # Общая информация
    total = len(matches_list)
    wins = sum(1 for m in matches_list if m.result == 'win')
    losses = total - wins
    winrate = round(wins / total * 100, 1) if total > 0 else 0

    info_para = doc.add_paragraph()
    info_para.add_run(f'Всего матчей: {total}  |  ').bold = False
    info_para.add_run(f'Побед: {wins}  |  ').bold = False
    info_para.add_run(f'Поражений: {losses}  |  ').bold = False
    run_wr = info_para.add_run(f'Винрейт: {winrate}%')
    run_wr.bold = True
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Таблица матчей
    headers = ['Дата', 'Герой', 'Тип', 'Результат', 'K/D/A', 'KDA',
               'Длит.', 'Роль', 'GPM', 'XPM']

    rows = []
    for m in matches_list:
        rows.append([
            m.date_played.strftime('%d.%m.%Y %H:%M'),
            m.hero.name,
            m.match_type_ref.name if m.match_type_ref else '—',
            'Победа' if m.result == 'win' else 'Поражение',
            m.kda_string,
            str(m.kda_ratio),
            f'{m.duration_minutes} мин',
            m.team_role,
            str(m.gpm),
            str(m.xpm)
        ])

    if rows:
        add_table_to_doc(doc, headers, rows)
    else:
        doc.add_paragraph('Нет записанных матчей.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    file_stream = save_doc_to_bytes(doc)
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'matches_{current_user.username}.docx'
    )


@app.route('/export/stats/word')
@login_required
def export_stats_word():
    doc = create_word_document(f'Статистика игрока — {current_user.username}')

    # --- Общая статистика ---
    doc.add_heading('Общая статистика', level=2)

    total_matches = current_user.matches.count()
    wins = current_user.matches.filter_by(result='win').count()
    losses = total_matches - wins
    winrate = current_user.get_winrate()
    avg_kda = current_user.get_avg_kda()
    best_hero = current_user.get_best_hero()

    # Средние
    avg_stats = db.session.query(
        db.func.avg(Match.kills).label('avg_kills'),
        db.func.avg(Match.deaths).label('avg_deaths'),
        db.func.avg(Match.assists).label('avg_assists'),
        db.func.avg(Match.duration_minutes).label('avg_duration'),
        db.func.avg(Match.gpm).label('avg_gpm'),
        db.func.avg(Match.xpm).label('avg_xpm')
    ).filter_by(user_id=current_user.id).first()

    general_headers = ['Показатель', 'Значение']
    general_rows = [
        ['Всего матчей', str(total_matches)],
        ['Побед', str(wins)],
        ['Поражений', str(losses)],
        ['Винрейт', f'{winrate}%'],
        ['Средний KDA', str(avg_kda)],
        ['Лучший герой', f'{best_hero["name"]} ({best_hero["winrate"]}%, {best_hero["games"]} игр)' if best_hero else '—'],
        ['Ср. убийства за матч', f'{avg_stats.avg_kills:.1f}' if avg_stats.avg_kills else '0'],
        ['Ср. смерти за матч', f'{avg_stats.avg_deaths:.1f}' if avg_stats.avg_deaths else '0'],
        ['Ср. помощь за матч', f'{avg_stats.avg_assists:.1f}' if avg_stats.avg_assists else '0'],
        ['Ср. длительность', f'{avg_stats.avg_duration:.0f} мин' if avg_stats.avg_duration else '0'],
        ['Ср. GPM', f'{avg_stats.avg_gpm:.0f}' if avg_stats.avg_gpm else '0'],
        ['Ср. XPM', f'{avg_stats.avg_xpm:.0f}' if avg_stats.avg_xpm else '0'],
    ]

    add_table_to_doc(doc, general_headers, general_rows)
    doc.add_paragraph()

    # --- Статистика по героям ---
    hero_stats = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins'),
        db.func.avg(Match.kills).label('avg_k'),
        db.func.avg(Match.deaths).label('avg_d'),
        db.func.avg(Match.assists).label('avg_a')
    ).join(Match).filter(
        Match.user_id == current_user.id
    ).group_by(Hero.id).order_by(db.text('games DESC')).all()

    if hero_stats:
        doc.add_heading('Статистика по героям', level=2)

        hero_headers = ['Герой', 'Игры', 'Победы', 'Винрейт', 'Ср. K', 'Ср. D', 'Ср. A']
        hero_rows = []
        for h in hero_stats:
            wr = round(h.wins / h.games * 100, 1) if h.games > 0 else 0
            hero_rows.append([
                h.name, str(h.games), str(h.wins), f'{wr}%',
                f'{h.avg_k:.1f}', f'{h.avg_d:.1f}', f'{h.avg_a:.1f}'
            ])

        add_table_to_doc(doc, hero_headers, hero_rows)
        doc.add_paragraph()

    # --- Статистика по ролям ---
    role_stats = db.session.query(
        Match.team_role,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).filter_by(user_id=current_user.id).group_by(Match.team_role).all()

    if role_stats:
        doc.add_heading('Статистика по ролям', level=2)

        role_headers = ['Роль', 'Игры', 'Победы', 'Винрейт']
        role_rows = []
        for r in role_stats:
            wr = round(r.wins / r.games * 100, 1) if r.games > 0 else 0
            role_rows.append([r.team_role, str(r.games), str(r.wins), f'{wr}%'])

        add_table_to_doc(doc, role_headers, role_rows)

    file_stream = save_doc_to_bytes(doc)
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'stats_{current_user.username}.docx'
    )


@app.route('/export/report/word')
@login_required
def export_report_word():
    doc = create_word_document(f'Аналитический отчёт — {current_user.username}')

    # --- Топ героев по винрейту ---
    top_heroes = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).join(Match).filter(
        Match.user_id == current_user.id
    ).group_by(Hero.id).having(
        db.func.count(Match.id) >= 1
    ).all()

    if top_heroes:
        doc.add_heading('Топ героев по винрейту', level=2)

        top_data = []
        for h in top_heroes:
            wr = round(h.wins / h.games * 100, 1) if h.games > 0 else 0
            top_data.append((h.name, h.games, h.wins, wr))

        top_data.sort(key=lambda x: x[3], reverse=True)

        top_headers = ['#', 'Герой', 'Игры', 'Победы', 'Винрейт']
        top_rows = []
        for i, (name, games, wins, wr) in enumerate(top_data[:15], 1):
            top_rows.append([str(i), name, str(games), str(wins), f'{wr}%'])

        add_table_to_doc(doc, top_headers, top_rows)
        doc.add_paragraph()

    # --- Статистика по ролям ---
    role_stats = db.session.query(
        Match.team_role,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins'),
        db.func.avg(Match.kills).label('avg_k'),
        db.func.avg(Match.deaths).label('avg_d'),
        db.func.avg(Match.assists).label('avg_a')
    ).filter_by(user_id=current_user.id).group_by(Match.team_role).all()

    if role_stats:
        doc.add_heading('Статистика по ролям в команде', level=2)

        role_headers = ['Роль', 'Игры', 'Победы', 'Винрейт', 'Ср. K', 'Ср. D', 'Ср. A']
        role_rows = []
        for r in role_stats:
            wr = round(r.wins / r.games * 100, 1) if r.games > 0 else 0
            role_rows.append([
                r.team_role, str(r.games), str(r.wins), f'{wr}%',
                f'{r.avg_k:.1f}', f'{r.avg_d:.1f}', f'{r.avg_a:.1f}'
            ])

        add_table_to_doc(doc, role_headers, role_rows)
        doc.add_paragraph()

    # --- Активность по месяцам ---
    monthly = db.session.query(
        db.func.strftime('%Y-%m', Match.date_played).label('month'),
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).filter_by(user_id=current_user.id).group_by('month').order_by('month').all()

    if monthly:
        doc.add_heading('Активность по месяцам', level=2)

        month_headers = ['Месяц', 'Игры', 'Победы', 'Винрейт']
        month_rows = []
        for m in monthly:
            wr = round(m.wins / m.games * 100, 1) if m.games > 0 else 0
            month_rows.append([m.month, str(m.games), str(m.wins), f'{wr}%'])

        add_table_to_doc(doc, month_headers, month_rows)
        doc.add_paragraph()

    # --- Итоговая сводка ---
    doc.add_heading('Итоговая сводка', level=2)

    total = current_user.matches.count()
    wins_count = current_user.matches.filter_by(result='win').count()
    winrate = current_user.get_winrate()
    avg_kda = current_user.get_avg_kda()
    best = current_user.get_best_hero()

    summary = doc.add_paragraph()
    summary.add_run(f'За всё время сыграно {total} матчей. ').font.size = Pt(12)
    summary.add_run(f'Из них {wins_count} побед и {total - wins_count} поражений. ').font.size = Pt(12)
    summary.add_run(f'Общий винрейт составляет {winrate}%. ').font.size = Pt(12)
    summary.add_run(f'Средний показатель KDA: {avg_kda}. ').font.size = Pt(12)
    if best:
        summary.add_run(
            f'Лучший герой — {best["name"]} с винрейтом {best["winrate"]}% за {best["games"]} игр.'
        ).font.size = Pt(12)

    file_stream = save_doc_to_bytes(doc)
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'report_{current_user.username}.docx'
    )

# ==========================================
#          АДМИНКА
# ==========================================

@app.route('/admin')
@login_required
@admin_required
def admin_dashboard():
    total_users = User.query.count()
    total_matches = Match.query.count()
    total_heroes = Hero.query.count()

    # Самые активные игроки
    active_players = db.session.query(
        User.username,
        db.func.count(Match.id).label('games')
    ).join(Match).group_by(User.id).order_by(db.text('games DESC')).limit(10).all()

    # Самые популярные герои
    popular_heroes = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('picks'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).join(Match).group_by(Hero.id).order_by(db.text('picks DESC')).limit(10).all()

    return render_template('admin/dashboard.html',
                           total_users=total_users,
                           total_matches=total_matches,
                           total_heroes=total_heroes,
                           active_players=active_players,
                           popular_heroes=popular_heroes)


# --- Админ: Пользователи ---

@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    search = request.args.get('search', '')
    query = User.query
    if search:
        query = query.filter(
            db.or_(User.username.ilike(f'%{search}%'),
                   User.email.ilike(f'%{search}%'))
        )
    users = query.order_by(User.created_at.desc()).all()
    return render_template('admin/users.html', users=users, search=search)


@app.route('/admin/users/<int:user_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_edit_user(user_id):
    user = User.query.get_or_404(user_id)
    form = UserEditForm(obj=user)
    if form.validate_on_submit():
        # Проверяем уникальность
        existing = User.query.filter(User.username == form.username.data,
                                     User.id != user.id).first()
        if existing:
            flash('Имя пользователя уже занято.', 'danger')
        else:
            user.username = form.username.data
            user.email = form.email.data
            user.is_admin = form.is_admin.data
            user.is_active_user = form.is_active_user.data
            db.session.commit()
            flash('Пользователь обновлён.', 'success')
            return redirect(url_for('admin_users'))
    return render_template('admin/users.html', users=User.query.all(),
                           edit_user=user, form=form, search='')


@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_user(user_id):
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        flash('Нельзя удалить самого себя.', 'danger')
        return redirect(url_for('admin_users'))
    db.session.delete(user)
    db.session.commit()
    flash('Пользователь удалён.', 'info')
    return redirect(url_for('admin_users'))


# --- Админ: Герои ---

@app.route('/admin/heroes')
@login_required
@admin_required
def admin_heroes():
    heroes_list = Hero.query.order_by(Hero.name).all()
    return render_template('admin/heroes.html', heroes=heroes_list)


@app.route('/admin/heroes/add', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_add_hero():
    form = HeroForm()
    form.role_id.choices = [(r.id, r.name) for r in HeroRole.query.order_by(HeroRole.name).all()]

    if form.validate_on_submit():
        hero = Hero(
            name=form.name.data,
            attribute=form.attribute.data,
            role_id=form.role_id.data,
            difficulty=form.difficulty.data,
            attack_type=form.attack_type.data,
            description=form.description.data or ''
        )
        db.session.add(hero)
        db.session.commit()
        flash('Герой добавлен!', 'success')
        return redirect(url_for('admin_heroes'))

    return render_template('admin/hero_form.html', form=form, title='Добавить героя')


@app.route('/admin/heroes/<int:hero_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_edit_hero(hero_id):
    hero = Hero.query.get_or_404(hero_id)
    form = HeroForm(obj=hero)
    form.role_id.choices = [(r.id, r.name) for r in HeroRole.query.order_by(HeroRole.name).all()]

    if form.validate_on_submit():
        hero.name = form.name.data
        hero.attribute = form.attribute.data
        hero.role_id = form.role_id.data
        hero.difficulty = form.difficulty.data
        hero.attack_type = form.attack_type.data
        hero.description = form.description.data or ''
        db.session.commit()
        flash('Герой обновлён!', 'success')
        return redirect(url_for('admin_heroes'))

    return render_template('admin/hero_form.html', form=form,
                           title='Редактировать героя', hero=hero)


@app.route('/admin/heroes/<int:hero_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_hero(hero_id):
    hero = Hero.query.get_or_404(hero_id)
    db.session.delete(hero)
    db.session.commit()
    flash('Герой удалён.', 'info')
    return redirect(url_for('admin_heroes'))


# --- Админ: Матчи ---

@app.route('/admin/matches')
@login_required
@admin_required
def admin_matches():
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')

    query = Match.query

    if search:
        query = query.join(User).filter(
            db.or_(User.username.ilike(f'%{search}%'))
        )

    pagination = query.order_by(Match.date_played.desc()).paginate(
        page=page, per_page=30, error_out=False)

    return render_template('admin/matches.html', pagination=pagination, search=search)


@app.route('/admin/matches/<int:match_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_match(match_id):
    match = Match.query.get_or_404(match_id)
    db.session.delete(match)
    db.session.commit()
    flash('Матч удалён.', 'info')
    return redirect(url_for('admin_matches'))


# --- Админ: Справочники ---

@app.route('/admin/roles', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_roles():
    form = HeroRoleForm()
    if form.validate_on_submit():
        role = HeroRole(name=form.name.data)
        db.session.add(role)
        db.session.commit()
        flash('Роль добавлена!', 'success')
        return redirect(url_for('admin_roles'))
    roles = HeroRole.query.order_by(HeroRole.name).all()
    return render_template('admin/roles.html', roles=roles, form=form)


@app.route('/admin/roles/<int:role_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_role(role_id):
    role = HeroRole.query.get_or_404(role_id)
    if role.heroes.count() > 0:
        flash('Нельзя удалить роль, к которой привязаны герои.', 'danger')
    else:
        db.session.delete(role)
        db.session.commit()
        flash('Роль удалена.', 'info')
    return redirect(url_for('admin_roles'))


@app.route('/admin/match-types', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_match_types():
    form = MatchTypeForm()
    if form.validate_on_submit():
        mt = MatchType(name=form.name.data)
        db.session.add(mt)
        db.session.commit()
        flash('Тип матча добавлен!', 'success')
        return redirect(url_for('admin_match_types'))
    types = MatchType.query.order_by(MatchType.name).all()
    return render_template('admin/match_types.html', types=types, form=form)


@app.route('/admin/match-types/<int:type_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_match_type(type_id):
    mt = MatchType.query.get_or_404(type_id)
    if mt.matches.count() > 0:
        flash('Нельзя удалить тип, к которому привязаны матчи.', 'danger')
    else:
        db.session.delete(mt)
        db.session.commit()
        flash('Тип матча удалён.', 'info')
    return redirect(url_for('admin_match_types'))


# --- Админ: экспорт всех данных ---

@app.route('/admin/export/all-matches')
@login_required
@admin_required
def admin_export_all_matches():
    matches_list = Match.query.order_by(Match.date_played.desc()).all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID', 'Игрок', 'Дата', 'Герой', 'Тип матча', 'Результат',
                     'K', 'D', 'A', 'KDA', 'Длительность', 'Роль', 'GPM', 'XPM'])

    for m in matches_list:
        writer.writerow([
            m.id, m.player.username,
            m.date_played.strftime('%Y-%m-%d %H:%M'),
            m.hero.name,
            m.match_type_ref.name if m.match_type_ref else '',
            m.result, m.kills, m.deaths, m.assists, m.kda_ratio,
            m.duration_minutes, m.team_role, m.gpm, m.xpm
        ])

    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment;filename=all_matches_export.csv'}
    )

@app.route('/admin/export/all-matches/word')
@login_required
@admin_required
def admin_export_all_matches_word():
    matches_list = Match.query.order_by(Match.date_played.desc()).all()

    doc = create_word_document('Полный отчёт по всем матчам')

    # Общая статистика
    total = len(matches_list)
    total_users = User.query.count()
    total_heroes = Hero.query.count()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Пользователей: {total_users}  |  Матчей: {total}  |  Героев: {total_heroes}')

    doc.add_paragraph()

    # Таблица всех матчей
    headers = ['ID', 'Игрок', 'Дата', 'Герой', 'Тип', 'Результат',
               'K/D/A', 'KDA', 'Длит.', 'Роль', 'GPM', 'XPM']

    rows = []
    for m in matches_list:
        rows.append([
            str(m.id),
            m.player.username,
            m.date_played.strftime('%d.%m.%Y %H:%M'),
            m.hero.name,
            m.match_type_ref.name if m.match_type_ref else '—',
            'Победа' if m.result == 'win' else 'Поражение',
            m.kda_string,
            str(m.kda_ratio),
            f'{m.duration_minutes}',
            m.team_role,
            str(m.gpm),
            str(m.xpm)
        ])

    if rows:
        add_table_to_doc(doc, headers, rows)
    else:
        doc.add_paragraph('Нет записанных матчей.')

    doc.add_paragraph()

    # Статистика по игрокам
    active_players = db.session.query(
        User.username,
        db.func.count(Match.id).label('games'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).join(Match).group_by(User.id).order_by(db.text('games DESC')).all()

    if active_players:
        doc.add_heading('Статистика по игрокам', level=2)

        player_headers = ['Игрок', 'Матчей', 'Побед', 'Винрейт']
        player_rows = []
        for p in active_players:
            wr = round(p.wins / p.games * 100, 1) if p.games > 0 else 0
            player_rows.append([p.username, str(p.games), str(p.wins), f'{wr}%'])

        add_table_to_doc(doc, player_headers, player_rows)
        doc.add_paragraph()

    # Статистика по героям (глобальная)
    hero_global = db.session.query(
        Hero.name,
        db.func.count(Match.id).label('picks'),
        db.func.sum(db.case((Match.result == 'win', 1), else_=0)).label('wins')
    ).join(Match).group_by(Hero.id).order_by(db.text('picks DESC')).limit(20).all()

    if hero_global:
        doc.add_heading('Топ героев (глобально)', level=2)

        hero_headers = ['Герой', 'Пиков', 'Побед', 'Винрейт']
        hero_rows = []
        for h in hero_global:
            wr = round(h.wins / h.picks * 100, 1) if h.picks > 0 else 0
            hero_rows.append([h.name, str(h.picks), str(h.wins), f'{wr}%'])

        add_table_to_doc(doc, hero_headers, hero_rows)

    file_stream = save_doc_to_bytes(doc)
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='all_matches_report.docx'
    )

# ==========================================
#     Создание БД и запуск
# ==========================================

def create_tables():
    with app.app_context():
        db.create_all()

        # Создаём админа по умолчанию если нет
        if not User.query.filter_by(is_admin=True).first():
            admin = User(username='admin', email='admin@dota2tracker.com',
                         is_admin=True)
            admin.set_password('admin123')
            db.session.add(admin)
            db.session.commit()
            print('Создан админ: admin / admin123')

        # Создаём базовые справочные данные
        if HeroRole.query.count() == 0:
            roles = ['Carry', 'Nuker', 'Initiator', 'Disabler', 'Support',
                     'Durable', 'Escape', 'Pusher']
            for r in roles:
                db.session.add(HeroRole(name=r))
            db.session.commit()
            print('Созданы роли героев')

        if MatchType.query.count() == 0:
            types = ['Ranked All Pick', 'Unranked All Pick', 'Turbo',
                     'Captain Mode', 'Random Draft', 'Single Draft']
            for t in types:
                db.session.add(MatchType(name=t))
            db.session.commit()
            print('Созданы типы матчей')

def create_word_document(title):
    """Создаёт документ Word с базовым оформлением"""
    doc = Document()

    # Стили шрифтов по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Заголовок документа
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Дата генерации
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'Дата формирования: {datetime.utcnow().strftime("%d.%m.%Y %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # пустая строка

    return doc


def add_table_to_doc(doc, headers, rows):
    """Добавляет стилизованную таблицу в документ"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

    # Данные
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'

    return table


def save_doc_to_bytes(doc):
    """Сохраняет документ в байтовый поток"""
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

if __name__ == '__main__':
    create_tables()
    app.run(debug=True, port=5000)